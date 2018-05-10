import requests
import os, shutil, re
from docx import Document
# 定义url和头信息
url = 'http://bj.58.com/dashanzi/chuzu/pn1/?ClickID=1'
headers = {'User-Agent':'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:59.0) Gecko/20100101 Firefox/59.0'}

if not os.path.isfile("58city"):
    shutil.rmtree("58city")
os.mkdir("58city")

document = Document()
document.add_heading('58同城北京大山子租房数据')
try:
    r = requests.get(url, headers=headers)
    # 获取数据
    data = r.content.decode("utf-8") 
    # 括号是危险字符，切记切记
    # 使用正则表达式提取数据
    roomTitle1 = r'target="_blank"  rel="nofollow" >\s*?                        (.*?)                    </a>'
    Titles = re.compile(roomTitle1).findall(data)
    roomTitle = 'onclick="clickLog\(.from=fcpc_zflist_gzcount.\)."\s*?                       target="_blank" >\s*?                        (.*?)                    </a>'
    Titles.extend(re.compile(roomTitle).findall(data))
    print("已爬取到" + str(len(Titles)) + "条标题信息")
    roomImg = r'<img\s                        lazy_src="(http://pic..58cdn.com.cn/.*?)"\s                        src="http://img.58cdn.com.cn/ui9/house/list/lazy_pic.png">\s                                    </a>'
    Imgs = re.compile(roomImg).findall(data)
    print("已爬取到" + str(len(Imgs)) + "条图片信息")
    i = 1
    # 遍历图片列表 下载图片
    for img in Imgs:
        r = requests.get(img)
        with open('58city/' + str(i) + '.jpg', 'wb') as file:
            file.write(r.content)
        i += 1
    roomType = '<p class="room">(.*?)                    &nbsp;&nbsp;&nbsp;&nbsp;.*?</p>'
    Types = re.compile(roomType).findall(data)
    print("已爬取到" + str(len(Types)) + "条户型信息")
    roomMoney = '<div class="money">\s                    <b>(.*?)</b>元/月'
    Money = re.compile(roomMoney).findall(data)
    print("已爬取到" + str(len(Money)) + "条钱的信息")
    # 使用zip压缩数据
    zippeds = zip(Titles, Types, Money)
    i = 1
    for zipped in zippeds:
        # 将数据写入文件
        p = document.add_paragraph('标题 :' + str(zipped[0]) +"\n户型 :" + str(zipped[1]) + '\n价格 :' + zipped[2])
        document.add_picture('58city/' + str(i) + ".jpg")
        i += 1
except requests.exceptions.RequestException as e:
    print(e)
# 保存文件
document.save('58同城数据.docx')