from urllib import request, error
import os, shutil, re
from docx import Document
# 定义url和头信息
url = 'http://maoyan.com/board/4'
headers = {'User-Agent':'Mozilla/5.0 (X11; Ubuntu; Linux x86_64; rv:59.0) Gecko/20100101 Firefox/59.0'}
if not os.path.isfile("猫眼电影"):
    shutil.rmtree("猫眼电影")
os.mkdir("猫眼电影")

# 新建立一个document文件, 此格式可以写入图片
document = Document()
document.add_heading('猫眼电影Top100评分最高的电影')
j = 1
try:
    for i in range(0, 100, 10):
        print("正在爬取第" + str(i) + "页。。。")
        # 爬取数据
        req = request.Request(url + "?offset=" + str(i), headers=headers, method= 'GET')
        res = request.urlopen(req)
        data = res.read().decode("utf-8")
        moiveId = r'<i class="board-index board-index-(.*?)"'
        Ids = re.compile(moiveId).findall(data)
        print("已爬取到" + str(len(Ids)) + "条Id")
        moiveTitle = r'<a href="/films/.*?" title="(.*?)" class="image-link"'
        Titles = re.compile(moiveTitle).findall(data)
        print("已爬取到" + str(len(Titles)) + "条Title")
        moivelImg = r'<img data-src="(.*?)" alt='
        Imgs = re.compile(moivelImg).findall(data)
        print("已爬取到" + str(len(Imgs)) + "条Img")
        for img in Imgs:
            request.urlretrieve(img, "猫眼电影/" + str(j) + ".jpg")
            j += 1
        moiveStar = r'<p class="star">\s*?(.*?)\s*?</p>'
        Stars = re.compile(moiveStar).findall(data)
        print("已爬取到" + str(len(Stars)) + "条Star")
        moiveTime = r'上映时间：(.*?)</p>'
        Times = re.compile(moiveTime).findall(data)
        print("已爬取到第" + str(len(Times)) + "条Time")
        moiveGrade = r'(\d\.)</i><i class="fraction">(\d)'
        Grades = re.compile(moiveGrade).findall(data)
        print("已爬取到" + str(len(Grades)) + "条Grade")
        zippeds = zip(Ids, Titles, Stars, Times, Grades)
        # 存入文件
        for zipped in zippeds:
            p = document.add_paragraph('No :' + str(zipped[0]) + '\n电影名 :' + str(zipped[1]) + "\n" + str(zipped[2]) + "\n上映时间 :" + str(zipped[3]) + '\n评分 :' + zipped[4][0] + zipped[4][1])
            document.add_picture('猫眼电影/' + zipped[0] + ".jpg")
# 储存url处理
except error.HTTPError as e:
    print(e.reason)
except error.URLError as e:
    print(e.reason)

document.save('猫眼电影.docx')